import os
import io
import json
import math
import time
import requests
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st
from datetime import datetime, timedelta
from typing import Optional
from dotenv import load_dotenv

# LLM client (Groq)
try:
    from groq import Groq
except Exception:
    Groq = None  # will handle missing package gracefully

# Reporting libs
try:
    from docx import Document
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
except Exception:
    Document = None
    SimpleDocTemplate = None
    Paragraph = None
    Spacer = None
    getSampleStyleSheet = None

# ---------------------------
# Load environment
# ---------------------------
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
groq_client = None
if GROQ_API_KEY and Groq is not None:
    try:
        groq_client = Groq(api_key=GROQ_API_KEY)
    except Exception as e:
        groq_client = None

# ---------------------------
# Page config & CSS (hover flip cards + centered card-style chat)
# ---------------------------
st.set_page_config(page_title="AI Drift Radar", page_icon="📡", layout="wide")

# Important: CSS includes hover flip cards (front/back) accessible on desktop via hover.
# Note: On mobile hover won't work — flip-cards remain clickable fallback via :active.
st.markdown("""
<style>
/* Sidebar gradient */
[data-testid="stSidebar"] { background: linear-gradient(180deg, #eef2ff, #fffaf0); }

/* Main container centered card */
.main-container { display:flex; justify-content:center; }
.app-card {
  width: 980px;
  background: #fff;
  padding: 20px;
  border-radius: 12px;
  border: 1px solid #eef0f6;
  box-shadow: 0 8px 28px rgba(20,20,50,0.06);
  margin: 18px 0;
}

/* Chat message cards */
.chat-user {
  background: #e9f3ff;
  padding: 14px;
  border-radius: 12px;
  margin: 10px 0;
  max-width: 78%;
  align-self: flex-end;
}
.chat-assistant {
  background: #fbf7ff;
  padding: 14px;
  border-radius: 12px;
  margin: 10px 0;
  max-width: 78%;
  align-self: flex-start;
  border-left: 4px solid #d6b8ff;
}
/* Header and text styles */
.header { 
  font-size: 28px; 
  font-weight: 700; 
  margin-bottom: 6px; 
}
.sub { 
  color: #555; 
  margin-bottom: 12px; 
}
.small { 
  font-size: 13px; 
  color: #666; 
}

/* Grid for flip-cards */
.flip-grid {
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); /* Automatically adjust the grid */
  gap: 20px;
  margin-top: 12px;
}

/* Flip-card style */
.flip-card {
  perspective: 1200px;
  width: 80%;  /* Let the card take full width of the grid container */
  height: 200px; /* Set a height suitable for your content */
}

/* Flip-card inner container */
.flip-card-inner {
  position: relative;
  width: 80%;
  height: 100%;
  transition: transform 0.6s;
  transform-style: preserve-3d;
}

/* Hover effect for flipping */
.flip-card:hover .flip-card-inner {
  transform: rotateY(180deg);
}

/* Fallback for touch devices */
.flip-card:active .flip-card-inner { 
  transform: rotateY(180deg); 
}

/* Front and back of the flip-card */
.flip-card-front, .flip-card-back {
  position: absolute;
  width: 100%;
  height: 100%;
  -webkit-backface-visibility: hidden;
  backface-visibility: hidden;
  border-radius: 10px;
  padding: 16px;
  box-shadow: 0 6px 18px rgba(20, 20, 50, 0.04);
}

/* Front of the card with gradient background */
.flip-card-front {
  background: linear-gradient(180deg, #ffffff, #f7f9ff);
}

/* Back of the card with different gradient */
.flip-card-back {
  background: linear-gradient(180deg, #fff8f0, #fff);
  transform: rotateY(180deg);
}

/* Button and download styles */
.stButton>button {
  background: linear-gradient(90deg, #e6e6ff, #fff8e6);
  border: 1px solid #d7c9ff;
  border-radius: 8px;
  color: #111;
  font-weight: 600;
}

/* Responsive layout adjustments */
@media (max-width: 900px) {
  .app-card { 
    width: 94%; 
    padding: 14px; 
  }
  .flip-card { 
    height: 150px; 
  }
}


/* buttons and download */
.stButton>button {
  background: linear-gradient(90deg,#e6e6ff,#fff8e6);
  border: 1px solid #d7c9ff;
  border-radius: 8px;
  color: #111;
  font-weight: 600;
}

/* responsive */
@media (max-width: 900px) {
  .app-card { width: 94%; padding: 14px; }
  .flip-card { height: 150px; }
}
</style>
""", unsafe_allow_html=True)

# ---------------------------
# Session state initialization
# ---------------------------
if "messages" not in st.session_state:
    st.session_state.messages = []  # {role,user/assistant,content,domain,time}
if "last_drift" not in st.session_state:
    st.session_state.last_drift = {}
if "last_explanation" not in st.session_state:
    st.session_state.last_explanation = ""
if "domain" not in st.session_state:
    st.session_state.domain = ""  # assistant will ask if empty
if "agents_enabled" not in st.session_state:
    st.session_state.agents_enabled = bool(groq_client)

# ---------------------------
# Domain fuzzy resolver
# ---------------------------
# Map many user inputs to canonical domain names
DOMAIN_CANONICAL = {
    "e-commerce": ["ecommerce", "e-commerce", "ecom", "online retail", "shopping", "retail"],
    "finance": ["finance", "banking", "payments", "transactions", "fin"],
    "healthcare": ["healthcare", "medical", "health", "clinic", "hospital"],
    "manufacturing": ["manufacturing", "factory", "industrial", "production"],
    "saas": ["saas", "software", "software-as-a-service", "web app"],
    "logistics": ["logistics", "delivery", "shipping", "transport"],
    "edtech": ["edtech", "education", "learning", "school", "university"],
    "retail-offline": ["retail-offline", "offline retail", "brick and mortar", "store"],
    "insurance": ["insurance", "claims", "insurer"],
    "energy-iot": ["energy", "iot", "sensor", "meter", "energy-iot"]
}


# ---------------------------
# SMART DOMAIN RESOLVER (DROP-IN)
# ---------------------------
def resolve_domain(text):
    if not text:
        return None
    txt = text.lower().strip()

    # common greetings / short non-domain words to ignore
    invalid = {"hi", "hello", "hey", "yo", "help", "start", "run", "run analysis", "thanks", "thx"}
    if txt in invalid:
        return None

    # custom domain detection: "custom: airline" or "custom - airline"
    if txt.startswith("custom:") or txt.startswith("custom -"):
        parts = txt.split(":", 1) if ":" in txt else txt.split("-", 1)
        if len(parts) > 1 and parts[1].strip():
            return parts[1].strip().title()
        return "Custom"

    # fuzzy mapping of common aliases -> canonical domain names
    mapping = {
        "E-commerce": ["ecom", "e-commerce", "e commerce", "online retail", "shopping", "retail"],
        "Finance": ["finance", "fin", "banking", "payments", "transactions", "fintech"],
        "Healthcare": ["healthcare", "medical", "health", "clinic", "hospital", "med"],
        "Manufacturing": ["manufacturing", "factory", "industrial", "production"],
        "SaaS": ["saas", "software", "software-as-a-service", "web app", "subscription"],
        "Logistics": ["logistics", "delivery", "shipping", "transport", "supply chain"],
        "EdTech": ["edtech", "education", "learning", "school", "university"],
        "Retail-Offline": ["retail-offline", "offline retail", "store", "brick and mortar"],
        "Insurance": ["insurance", "claims", "insurer", "policy"],
        "Energy-IoT": ["energy", "iot", "meter", "smart meter", "power", "grid"]
    }

    # exact or alias match
    for canonical, aliases in mapping.items():
        if txt == canonical.lower() or txt in aliases:
            return canonical
        for a in aliases:
            if a in txt:
                return canonical

    # substring match for canonical with spaces
    for canonical in mapping:
        if canonical.replace("-", " ").lower() in txt:
            return canonical

    # fallback: short single-word input -> title-cased domain
    if len(txt.split()) <= 3:
        return txt.title()

    return None


# ---------------------------
# Core drift metric functions
# ---------------------------
def compute_psi_for_column(ref_series: pd.Series, cur_series: pd.Series, buckets: int = 10):
    try:
        a = ref_series.dropna().astype(float).values
        b = cur_series.dropna().astype(float).values
        if len(a) < 2 or len(b) < 2:
            return 0.0
        if len(a) > 5000:
            a = np.random.choice(a, 5000, replace=False)
        if len(b) > 5000:
            b = np.random.choice(b, 5000, replace=False)
        return psi(a, b, buckets=buckets)
    except Exception:
        return None


def compute_categorical_delta_for_column(ref_series: pd.Series, cur_series: pd.Series):
    try:
        return categorical_delta(ref_series.fillna(""), cur_series.fillna(""))
    except Exception:
        return None


# ---------------------------
# Embedding helpers
# ---------------------------
def mean_cosine_embedding_shift(ref_emb: np.ndarray, cur_emb: np.ndarray) -> float:
    try:
        if ref_emb is None or cur_emb is None:
            return 0.0
        ref_mean = np.mean(ref_emb, axis=0)
        cur_mean = np.mean(cur_emb, axis=0)
        denom = np.linalg.norm(ref_mean) * np.linalg.norm(cur_mean)
        if denom == 0:
            return 0.0
        cos_sim = float(np.dot(ref_mean, cur_mean) / denom)
        return max(0.0, 1.0 - cos_sim)
    except Exception:
        return 0.0


# ---------------------------
# Groq safe helpers (token extraction & sync)
# ---------------------------
SYSTEM_PROMPT_TEMPLATE = """
You are AI Drift Radar — an assistant for drift detection, model monitoring and integrations.
Context: domain = {domain}
Rules:
- Only answer drift, monitoring, instructions, or integration questions.
- If domain unspecified, ask user which domain to focus on.
Tone: helpful, concise, actionable.
"""


def _extract_token(chunk):
    try:
        choice = chunk.choices[0]
        delta = getattr(choice, "delta", None)
        if delta and getattr(delta, "content", None):
            return delta.content
        if isinstance(delta, dict):
            return delta.get("content", "") or ""
    except Exception:
        pass
    return ""


def stream_groq_answer(user_msg: str, domain: str, placeholder) -> str:
    if groq_client is None:
        return "Assistant disabled (no GROQ key)."
    system_prompt = SYSTEM_PROMPT_TEMPLATE.format(domain=domain or "unspecified")
    try:
        stream = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            temperature=0.2,
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_msg}],
            stream=True
        )
    except Exception as e:
        return f"Groq API error: {e}"
    full = ""
    for chunk in stream:
        token = _extract_token(chunk)
        if token:
            full += token
            placeholder.markdown(full)
    return full


def groq_complete_sync(prompt: str, domain: str, temperature: float = 0.2) -> str:
    if groq_client is None:
        return "Assistant disabled (no GROQ key)."
    system_prompt = SYSTEM_PROMPT_TEMPLATE.format(domain=domain or "unspecified")
    try:
        res = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            temperature=temperature,
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}],
            stream=False
        )
        # safest access
        try:
            return res.choices[0].message.content.strip()
        except Exception:
            try:
                return str(res.choices[0].message).strip()
            except Exception:
                return str(res)
    except Exception as e:
        return f"Groq API error: {e}"


# ---------------------------
# Agent shells (sync wrappers use groq_complete_sync)
# ---------------------------
def agent_drift_analyst(summary: str, domain: str) -> str:
    prompt = f"You are the Drift Analyst agent. Input summary:\\n{summary}\\nTask: Explain top drift features, quick checks, and numeric top-3."
    return groq_complete_sync(prompt, domain)


def agent_data_quality(ref_sample: str, cur_sample: str, domain: str) -> str:
    prompt = f"You are the Data Quality agent. Ref sample:\\n{ref_sample}\\nCur sample:\\n{cur_sample}\\nTask: List quality issues and quick fixes."
    return groq_complete_sync(prompt, domain)


def agent_business_impact(summary: str, domain: str) -> str:
    prompt = f"You are the Business Impact agent. Summary:\\n{summary}\\nTask: Explain business impact and priority."
    return groq_complete_sync(prompt, domain)


def agent_retrain_advisor(summary: str, domain: str) -> str:
    prompt = f"You are the Retrain Advisor agent. Summary:\\n{summary}\\nTask: Recommend retraining strategy and exact next steps (commands/pseudocode)."
    return groq_complete_sync(prompt, domain)


def agent_ops_integration(summary: str, domain: str) -> str:
    prompt = f"You are the Ops Integration agent. Summary:\\n{summary}\\nTask: Provide webhook payload, cURL, and monitoring checklist."
    return groq_complete_sync(prompt, domain)


# ---------------------------
# Reporting helpers (TXT/DOCX/PDF)
# ---------------------------
def make_txt(drift_scores: dict, explanation: str) -> bytes:
    lines = ["AI Drift Radar Report", f"Generated: {datetime.utcnow().isoformat()} UTC", "", "Drift Scores:"]
    for k, v in drift_scores.items():
        lines.append(f"{k}: {v}")
    lines.extend(["", "Explanation:", explanation])
    return "\n".join(lines).encode("utf-8")


def make_docx(drift_scores: dict, explanation: str) -> io.BytesIO:
    if Document is None:
        raise RuntimeError("python-docx not installed")
    doc = Document()
    doc.add_heading("AI Drift Radar Report", level=1)
    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()} UTC")
    doc.add_heading("Drift Scores", level=2)
    for k, v in drift_scores.items():
        doc.add_paragraph(f"{k}: {v}")
    doc.add_heading("Explanation", level=2)
    doc.add_paragraph(explanation)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def make_pdf(drift_scores: dict, explanation: str) -> io.BytesIO:
    if SimpleDocTemplate is None:
        raise RuntimeError("reportlab not installed")
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio)
    styles = getSampleStyleSheet()
    story = [Paragraph("AI Drift Radar Report", styles["Title"]), Spacer(1, 8),
             Paragraph(f"Generated: {datetime.utcnow().isoformat()} UTC", styles["Normal"]), Spacer(1, 12),
             Paragraph("Drift Scores:", styles["Heading2"])]
    for k, v in drift_scores.items():
        story.append(Paragraph(f"{k}: {v}", styles["Normal"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Explanation:", styles["Heading2"]))
    for para in explanation.split("\n\n"):
        story.append(Paragraph(para.replace("\n", "<br/>"), styles["Normal"]))
        story.append(Spacer(1, 6))
    doc.build(story)
    bio.seek(0)
    return bio


# ---------------------------
# Sample generator helpers (returns two CSV bytes)
# ---------------------------
DOMAINS = [
    "E-commerce", "Finance", "Healthcare", "Manufacturing", "SaaS",
    "Logistics", "EdTech", "Retail-Offline", "Insurance", "Energy-IoT"
]


def generate_sample_pair(domain: str, n_rows: int = 20000, seasonal_keyword: Optional[str] = None):
    rng = np.random.default_rng(12345)
    base_time = datetime.utcnow() - timedelta(days=90)
    ref_rows, cur_rows = [], []
    for i in range(n_rows):
        ts_ref = base_time + timedelta(minutes=int(rng.integers(0, 60 * 24 * 60)))
        ts_cur = datetime.utcnow() - timedelta(minutes=int(rng.integers(0, 60 * 24 * 7)))
        if domain.lower().startswith("e"):
            cats = ["Mobile", "Home", "Fashion", "Grocery", "Books"]
            cat_ref = rng.choice(cats)
            query_ref = rng.choice(["best price", "buy online", "top rated"])
            purchased_ref = int(rng.random() < 0.08)
            if seasonal_keyword and rng.random() < 0.45:
                cat_cur = seasonal_keyword + " Specials"
                query_cur = f"{seasonal_keyword} {rng.choice(['sale', 'offers', 'gift'])}"
                purchased_cur = int(rng.random() < 0.30)
            else:
                cat_cur = rng.choice(cats)
                query_cur = rng.choice(["best price", "discount"])
                purchased_cur = int(rng.random() < 0.09)
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "category": cat_ref, "query": query_ref,
                             "session_sec": int(abs(rng.normal(180, 60))), "purchased": purchased_ref})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "category": cat_cur, "query": query_cur,
                             "session_sec": int(abs(rng.normal(220, 80))), "purchased": purchased_cur})
        elif domain.lower().startswith("f"):
            # finance
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "acct_age_days": int(abs(rng.normal(400, 250))),
                             "tx_amount": round(abs(rng.normal(150, 400)), 2),
                             "tx_type": rng.choice(["payment", "transfer"]), "is_fraud": int(rng.random() < 0.01)})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "acct_age_days": int(abs(rng.normal(380, 260))),
                             "tx_amount": round(abs(rng.normal(160, 500)), 2),
                             "tx_type": rng.choice(["payment", "transfer", "refund"]),
                             "is_fraud": int(rng.random() < 0.012)})
        elif domain.lower().startswith("h"):
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "age": int(abs(rng.normal(50, 18))),
                             "glucose": round(abs(rng.normal(95, 18)), 1), "wbc": round(abs(rng.normal(6, 1.5)), 2)})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "age": int(abs(rng.normal(51, 19))),
                             "glucose": round(abs(rng.normal(100, 25)), 1), "wbc": round(abs(rng.normal(6.2, 1.6)), 2)})
        elif domain.lower().startswith("m"):
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "sensor_temp": round(50 + rng.normal(0, 4), 2),
                             "vibration": round(abs(rng.normal(0.3, 0.08)), 3),
                             "status": rng.choice(["ok", "warning"])})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "sensor_temp": round(50 + rng.normal(1, 5), 2),
                             "vibration": round(abs(rng.normal(0.45, 0.12)), 3),
                             "status": rng.choice(["ok", "warning", "fail"])})
        elif domain.lower().startswith("s"):
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "user_id": rng.integers(1000, 9999),
                             "active_sec": int(abs(rng.normal(600, 300))), "events": int(abs(rng.normal(10, 6)))})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "user_id": rng.integers(1000, 9999),
                             "active_sec": int(abs(rng.normal(700, 350))), "events": int(abs(rng.normal(14, 8)))})
        elif domain.lower().startswith("l"):
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "route": rng.integers(1, 100),
                             "duration_min": int(abs(rng.normal(50, 30))), "delay": int(abs(rng.normal(5, 10)))})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "route": rng.integers(1, 100),
                             "duration_min": int(abs(rng.normal(60, 40))), "delay": int(abs(rng.normal(10, 20)))})
        elif domain.lower().startswith("e") and "ed" in domain.lower():
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "student_id": rng.integers(1000, 9999),
                             "time_min": int(abs(rng.normal(30, 20))), "completed": int(rng.random() < 0.12)})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "student_id": rng.integers(1000, 9999),
                             "time_min": int(abs(rng.normal(40, 25))), "completed": int(rng.random() < 0.08)})
        elif domain.lower().startswith("r"):
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "store_id": rng.integers(1, 200),
                             "footfall": int(abs(rng.normal(120, 60))), "sales": round(abs(rng.normal(2000, 1500)), 2)})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "store_id": rng.integers(1, 200),
                             "footfall": int(abs(rng.normal(150, 80))), "sales": round(abs(rng.normal(2500, 1600)), 2)})
        elif domain.lower().startswith("i"):
            ref_rows.append(
                {"timestamp": ts_ref.isoformat(sep=' '), "claim_amount": round(abs(rng.normal(4000, 1800)), 2),
                 "claim_type": rng.choice(["auto", "health"]), "fraud_score": round(rng.random(), 3)})
            cur_rows.append(
                {"timestamp": ts_cur.isoformat(sep=' '), "claim_amount": round(abs(rng.normal(4200, 2000)), 2),
                 "claim_type": rng.choice(["auto", "health", "property"]), "fraud_score": round(rng.random(), 3)})
        elif domain.lower().startswith("en"):
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "meter": rng.integers(100, 999),
                             "power_kw": round(abs(rng.normal(5, 1.5)), 3),
                             "voltage": round(220 + rng.normal(0, 4), 2)})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "meter": rng.integers(100, 999),
                             "power_kw": round(abs(rng.normal(6, 2)), 3), "voltage": round(220 + rng.normal(1, 6), 2)})
        else:
            # default simple rows
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "value": round(abs(rng.normal(100, 40)), 2)})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "value": round(abs(rng.normal(120, 60)), 2)})
    df_ref = pd.DataFrame(ref_rows)
    df_cur = pd.DataFrame(cur_rows)
    return df_ref.to_csv(index=False).encode("utf-8"), df_cur.to_csv(index=False).encode("utf-8")


# ---------------------------
# Sample metrics.json + embeddings generator
# ---------------------------
def generate_sample_metrics_json(domain: str):
    # realistic defaults vary by domain; keep it generic
    sample = {"f1": 0.78, "roc": 0.84, "precision": 0.75, "recall": 0.72}
    return json.dumps(sample, indent=2).encode("utf-8")


def generate_sample_embeddings_pair(n_samples: int, dim: int, shift: float = 0.8):
    rng = np.random.default_rng(42)
    ref = rng.normal(0, 1, (n_samples, dim)).astype(np.float32)
    cur = (rng.normal(shift, 1, (n_samples, dim))).astype(np.float32)
    # return as bytes using np.save to buffer
    buf_ref = io.BytesIO()
    np.save(buf_ref, ref)
    buf_ref.seek(0)
    buf_cur = io.BytesIO()
    np.save(buf_cur, cur)
    buf_cur.seek(0)
    return buf_ref.read(), buf_cur.read()



# -----------------------------------------------
# Sidebar Navigation
# -----------------------------------------------
PAGES = [
    "Home",
    "FAQs",
    "Instructions",
    "Sample Data",
    "Upload & Analyze",
    "Model Monitor",
    "AI Assistant",
    "About"
]

st.sidebar.title("📡 AI Drift Radar")
page = st.sidebar.radio("Navigation", PAGES)

# -----------------------------------------------
# HOME PAGE
# -----------------------------------------------
if page == "Home":
    st.markdown("""
    <div class='header'>📡 AI Drift Radar</div>
    <div class='sub'>
        Your intelligent companion for detecting data drift, model degradation, and guiding retraining decisions —
        with real-time insights powered by Groq Llama 3.1.
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    ### ✨ What This System Does
    - Compares **old vs current data**  
    - Detects **feature drift** (numeric PSI & categorical)
    - Detects **embedding drift**  
    - Reads your **model performance metrics**  
    - Gives **business & technical explanation**  
    - Recommends **when to retrain**  
    - Generates **PDF / DOCX / TXT reports** 
    - Includes **AI agents**: Drift Analyst, Business Impact, Data Quality & more  
    - Lets you **generate sample data & embeddings** for testing  
    """)

    st.markdown("</div></div>", unsafe_allow_html=True)

elif page == "FAQs":
    st.markdown("## 📘 Frequently Asked Questions")
    st.write("Click or hover on the cards below to reveal explanations.")

    # Flashcards for FAQ
    flashcards = [
        ("What is Drift?",
         "When new incoming data looks different from past data, the model becomes outdated and inaccurate."),
        ("Reference vs Current Data?",
         "Reference = old stable data. Current = recent live data. Drift is measured between these two."),
        ("Model Metrics (F1/ROC)?",
         "Numbers that show model quality. Falling metrics = model struggling due to drift or noise."),
        ("What Are Embeddings?",
         "A numerical representation of meaning (text, image, product). Drift = meaning patterns shifted."),
        ("Why Drift Matters?",
         "Real world changes — seasons, trends, new fraud patterns — break static models."),
        ("What This App Does?",
         "Detects drift, explains root cause, evaluates metrics & embeddings, and recommends retraining."),
        ("What Are Agents?",
         "Mini-AI specialists: Drift Analyst, Business Impact, Data Quality, Retrain Advisor & Ops Agent."),
        ("How Do Agents Help?",
         "They analyze your data, metrics & embeddings and provide clear actionable recommendations.")
    ]

    # Render flip cards
    st.markdown("<div class='flip-grid'>", unsafe_allow_html=True)
    for front, back in flashcards:
        st.markdown(f"""
        <div class='flip-card'>
            <div class='flip-card-inner'>
                <div class='flip-card-front'>
                    <b>{front}</b>
                </div>
                <div class='flip-card-back'>
                    {back}
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

# -----------------------------------------------
# INSTRUCTIONS — HOVER FLIP CARDS
# -----------------------------------------------
elif page == "Instructions":

    st.markdown("## 📘 Instructions")


    # -------------------------
    # Developer Steps
    # -------------------------
    st.markdown("## 🛠 How to Use This System ")

    st.markdown("""
    **Step 1 — Download Sample Data**  
    Go to *Sample Data* page → choose domain → download:
    - `reference_data.csv`
    - `current_data.csv`
    - `metrics.json` (optional)
    - `ref_embeddings.npy` (optional)
    - `cur_embeddings.npy` (optional)

    **Step 2 — Upload CSVs in “Upload & Analyze”**  
    Compute PSI + categorical drift + see radar chart.

    **Step 3 — Upload Metrics/Embeddings (optional)**  
    In “Model Monitor” page, drop:
    - metrics.json  
    - embeddings (.npy)

    **Step 4 — Trigger Agents**  
    Go to *AI Assistant*  
    - type “run analysis”  
    - agents combine drift + metrics + embeddings  
    - gives business + ML + ops + retrain recommendations

    **Step 5 — Download Reports**  
    Export findings as:
    - PDF  
    - DOCX  
    - TXT  
    """)

    # -------------------------
    # Code Snippets for Developers
    # -------------------------
    st.markdown("## 👨‍💻 How to Generate Metrics (metrics.json)")
    st.code("""
from sklearn.metrics import f1_score, roc_auc_score, precision_score, recall_score
import json

metrics = {
    "f1": float(f1_score(y_true, y_pred)),
    "roc": float(roc_auc_score(y_true, y_prob)),
    "precision": float(precision_score(y_true, y_pred)),
    "recall": float(recall_score(y_true, y_pred))
}

with open("metrics.json", "w") as f:
    json.dump(metrics, f, indent=2)
""")

    st.markdown("## 👨‍💻 How to Generate Embeddings (.npy)")
    st.code("""
import numpy as np

ref_emb = model.encode(reference_inputs)
cur_emb = model.encode(current_inputs)

np.save("ref_embeddings.npy", ref_emb)
np.save("cur_embeddings.npy", cur_emb)
""")

    st.markdown("</div></div>", unsafe_allow_html=True)

# -----------------------------------------------
# SAMPLE DATA PAGE — CSV + METRICS + EMBEDDINGS GENERATOR
# -----------------------------------------------
elif page == "Sample Data":

    st.markdown("## 🧪 Sample Data Generator")
    st.write("Generate test-ready data for any domain.")

    domain = st.selectbox("Select Domain", DOMAINS)

    seasonal = st.text_input("Optional Seasonal Keyword (e.g., Christmas, Summer, Sale)")
    nrows = st.slider("Number of rows", 2000, 50000, 20000)

    if st.button("Generate Reference & Current CSV"):
        ref_csv, cur_csv = generate_sample_pair(domain, n_rows=nrows, seasonal_keyword=seasonal or None)

        col1, col2 = st.columns(2)
        with col1:
            st.download_button("Download reference_data.csv",
                               ref_csv, "reference_data.csv", "text/csv")
        with col2:
            st.download_button("Download current_data.csv",
                               cur_csv, "current_data.csv", "text/csv")

    st.markdown("---")

    st.markdown("### 📊 Generate metrics.json")
    if st.button("Generate Metrics File"):
        metrics_bytes = generate_sample_metrics_json(domain)
        st.download_button("Download metrics.json",
                           metrics_bytes, "metrics.json", "application/json")

    st.markdown("---")

    st.markdown("### 🔢 Generate Embeddings (.npy)")

    emb_dim = st.selectbox("Embedding dimension", [32, 64, 128, 256], index=1)

    # Embedding dimension explanation:
    dim_explain = {
        32: "Useful for tiny recommendation systems, lightweight mobile ML.",
        64: "Balanced — good for most text/product embeddings.",
        128: "Higher precision, better semantic separation.",
        256: "Best for rich semantic domains (vision, multimodal, deep ranking)."
    }
    st.info(dim_explain[emb_dim])

    if st.button("Generate Embeddings"):
        ref_emb, cur_emb = generate_sample_embeddings_pair(2000, dim=emb_dim)

        col1, col2 = st.columns(2)
        with col1:
            st.download_button("Download ref_embeddings.npy", ref_emb,
                               "ref_embeddings.npy", "application/octet-stream")
        with col2:
            st.download_button("Download cur_embeddings.npy", cur_emb,
                               "cur_embeddings.npy", "application/octet-stream")

    st.markdown("</div></div>", unsafe_allow_html=True)

## ---------------------------
# PAGE: Upload & Analyze (ROBUST DRIFT)
# Replace existing Upload & Analyze page with this block
# ---------------------------
import numpy as np
from scipy.stats import ks_2samp


import numpy as np
from scipy.stats import ks_2samp
import plotly.graph_objects as go


# ---------------------------
# CLEAN PREPROCESSOR
# ---------------------------
def preprocess_for_drift(df):
    df = df.copy()

    # Fix numeric-like strings ("1,234", "20", "0")
    for col in df.columns:
        if df[col].dtype == object:
            df[col] = (
                df[col].astype(str)
                .str.replace(",", "", regex=False)
                .str.strip()
            )
            df[col] = pd.to_numeric(df[col], errors="ignore")

    # Expand timestamp if present
    for col in df.columns:
        low = col.lower()
        if any(k in low for k in ["timestamp", "date", "time"]):
            try:
                df[col] = pd.to_datetime(df[col], errors="coerce")
                df[col + "_hour"] = df[col].dt.hour
                df[col + "_day"] = df[col].dt.day
                df[col + "_weekday"] = df[col].dt.weekday
            except:
                pass

    return df


# ---------------------------
# COLUMN TYPE DETECTOR
# ---------------------------
def identify_valid_columns(df):
    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()

    categorical_cols = []
    for col in df.select_dtypes(include=["object"]).columns:
        unique_count = df[col].nunique()
        if 1 < unique_count <= 30:  # treat only *low-cardinality* text as categorical
            categorical_cols.append(col)

    return numeric_cols, categorical_cols


# ---------------------------
# NUMERIC DRIFT (KS Test)
# ---------------------------
def compute_numeric_drift(ref, cur):
    results = {}
    for col in ref.columns:
        r, c = ref[col].dropna(), cur[col].dropna()

        if len(r) < 3 or len(c) < 3:
            continue

        try:
            stat, _ = ks_2samp(r, c)
            results[col] = float(stat)
        except:
            pass

    return results


# ---------------------------
# CATEGORICAL DRIFT (PSI)
# ---------------------------
def compute_categorical_drift(ref, cur):
    results = {}

    for col in ref.columns:
        r_counts = ref[col].astype(str).value_counts(normalize=True)
        c_counts = cur[col].astype(str).value_counts(normalize=True)

        categories = set(r_counts.index) | set(c_counts.index)
        psi = 0.0

        for cat in categories:
            r_p = r_counts.get(cat, 1e-6)
            c_p = c_counts.get(cat, 1e-6)
            psi += (r_p - c_p) * np.log((r_p + 1e-9) / (c_p + 1e-9))

        results[col] = abs(float(psi))

    return results


# ---------------------------
# FULL DRIFT WRAPPER
# ---------------------------
def compute_full_drift(ref_df, cur_df):
    ref = preprocess_for_drift(ref_df)
    cur = preprocess_for_drift(cur_df)

    shared = list(set(ref.columns) & set(cur.columns))
    if not shared:
        return None, "No shared columns between reference and current data."

    ref = ref[shared]
    cur = cur[shared]

    ref_num, ref_cat = identify_valid_columns(ref)
    cur_num, cur_cat = identify_valid_columns(cur)

    numeric_cols = list(set(ref_num) & set(cur_num))
    categorical_cols = list(set(ref_cat) & set(cur_cat))

    ignored_cols = [c for c in shared if c not in numeric_cols + categorical_cols]

    drift = {}

    if numeric_cols:
        drift.update(compute_numeric_drift(ref[numeric_cols], cur[numeric_cols]))

    if categorical_cols:
        drift.update(compute_categorical_drift(ref[categorical_cols], cur[categorical_cols]))

    if not drift:
        return None, f"Could not compute drift. Usable = {numeric_cols + categorical_cols}; Ignored = {ignored_cols}"

    return {
        "drift": drift,
        "numeric_columns": numeric_cols,
        "categorical_columns": categorical_cols,
        "ignored_columns": ignored_cols
    }, None


## ============================================
# UPLOAD & ANALYZE — FINAL FULL BLOCK (DROP-IN)
# ============================================

import numpy as np
from scipy.stats import ks_2samp
import plotly.graph_objects as go
from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet


# ------------------------------------------------------
# ===== EXPORT HELPERS: TXT / DOCX / PDF
# ------------------------------------------------------
def export_txt(text: str):
    return BytesIO(text.encode("utf-8"))


def export_docx(text: str):
    buffer = BytesIO()
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_pdf(text: str):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    flow = []
    for line in text.split("\n"):
        flow.append(Paragraph(line, styles["Normal"]))
    doc.build(flow)
    buffer.seek(0)
    return buffer


# ------------------------------------------------------
# DOMAIN INFERENCE FROM DATA / DRIFT
# ------------------------------------------------------
def infer_domain_from_data(df=None, drift=None, metrics=None):
    # Check if df is a DataFrame (use it for domain inference)
    if df is not None:
        # Ensure the input is a DataFrame and get the column names
        cols = [c.lower() for c in df.columns]

        # Domain inference based on column names
        ec = ["session", "product", "category", "cart", "order", "sku"]
        fi = ["amount", "transaction", "balance", "loan", "credit"]
        hc = ["glucose", "wbc", "pulse", "patient"]
        mf = ["machine", "sensor", "pressure", "vibration", "rpm"]
        rt = ["footfall", "store", "region", "sales"]
        ss = ["tenant_id", "subscription", "user_id"]

        # Helper function to check column matches
        def match(keys):
            return any(k in col for col in cols for k in keys)

        if match(ec): return "E-commerce"
        if match(fi): return "Finance"
        if match(hc): return "Healthcare"
        if match(mf): return "Manufacturing"
        if match(rt): return "Retail-Offline"
        if match(ss): return "SaaS"

    # Check if drift is a dictionary (use its keys for domain inference)
    elif drift is not None and isinstance(drift, dict):
        dcols = [c.lower() for c in drift.keys()]

        # Domain inference based on drift data (dictionary)
        if "glucose" in dcols or "wbc" in dcols:
            return "Healthcare"
        if "session" in dcols or "category" in dcols:
            return "E-commerce"
        if "amount" in dcols or "transaction" in dcols:
            return "Finance"

    # Metric-based inference
    if metrics is not None:
        if "roc" in metrics or "precision" in metrics:
            return "Generic ML"

    # Return None if no inference matches
    return None

# ------------------------------------------------------
# DRIFT ENGINE (final stable)
# ------------------------------------------------------
def preprocess_for_drift(df):
    df = df.copy()

    for col in df.columns:
        if df[col].dtype == object:
            df[col] = (
                df[col].astype(str)
                .str.replace(",", "", regex=False)
                .str.strip()
            )
            df[col] = pd.to_numeric(df[col], errors="ignore")

    for col in df.columns:
        low = col.lower()
        if any(k in low for k in ["timestamp", "date", "time"]):
            try:
                df[col] = pd.to_datetime(df[col], errors="coerce")
                df[col + "_hour"] = df[col].dt.hour
                df[col + "_day"] = df[col].dt.day
                df[col + "_weekday"] = df[col].dt.weekday
            except:
                pass

    return df


def identify_valid_columns(df):
    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()

    categorical_cols = []
    for col in df.select_dtypes(include=['object']).columns:
        nu = df[col].nunique()
        if 1 < nu <= 30:
            categorical_cols.append(col)

    return numeric_cols, categorical_cols


def compute_numeric_drift(ref, cur):
    out = {}
    for col in ref.columns:
        r, c = ref[col].dropna(), cur[col].dropna()
        if len(r) < 3 or len(c) < 3:
            continue
        try:
            stat, _ = ks_2samp(r, c)
            out[col] = float(stat)
        except:
            pass
    return out


def compute_categorical_drift(ref, cur):
    out = {}
    for col in ref.columns:
        r_counts = ref[col].astype(str).value_counts(normalize=True)
        c_counts = cur[col].astype(str).value_counts(normalize=True)
        cats = set(r_counts.index) | set(c_counts.index)
        psi = 0.0
        for cat in cats:
            r_p = r_counts.get(cat, 1e-6)
            c_p = c_counts.get(cat, 1e-6)
            psi += (r_p - c_p) * np.log((r_p + 1e-9) / (c_p + 1e-9))
        out[col] = abs(float(psi))
    return out


def compute_full_drift(ref_df, cur_df):
    ref = preprocess_for_drift(ref_df)
    cur = preprocess_for_drift(cur_df)

    shared = list(set(ref.columns) & set(cur.columns))
    if not shared:
        return None, "No shared columns found."

    ref = ref[shared]
    cur = cur[shared]

    ref_num, ref_cat = identify_valid_columns(ref)
    cur_num, cur_cat = identify_valid_columns(cur)

    numeric_cols = list(set(ref_num) & set(cur_num))
    categorical_cols = list(set(ref_cat) & set(cur_cat))

    ignored = [c for c in shared if c not in numeric_cols + categorical_cols]

    drift = {}
    if numeric_cols:
        drift.update(compute_numeric_drift(ref[numeric_cols], cur[numeric_cols]))
    if categorical_cols:
        drift.update(compute_categorical_drift(ref[categorical_cols], cur[categorical_cols]))

    if not drift:
        return None, f"Could not compute drift. Usable={numeric_cols + categorical_cols}, Ignored={ignored}"

    return {
        "drift": drift,
        "numeric_columns": numeric_cols,
        "categorical_columns": categorical_cols,
        "ignored_columns": ignored,
        "ref_processed": ref,
        "cur_processed": cur
    }, None


# ------------------------------------------------------
# ===== STREAMLIT PAGE: Upload & Analyze
# ------------------------------------------------------
if page == "Upload & Analyze":
    st.markdown("## 📤 Upload & Analyze")

    ref_file = st.file_uploader("Upload Reference Dataset (CSV)", type=["csv"])
    cur_file = st.file_uploader("Upload Current Dataset (CSV)", type=["csv"])

    if ref_file and cur_file:

        df_ref = pd.read_csv(ref_file)
        df_cur = pd.read_csv(cur_file)

        st.success(f"Loaded {len(df_ref)} reference & {len(df_cur)} current rows.")

        drift_result, drift_err = compute_full_drift(df_ref, df_cur)

        if drift_err:
            st.error(drift_err)
            st.stop()

        st.json(drift_result)

        # Save drift for assistant
        st.session_state["last_drift"] = drift_result["drift"]

        # ---------------------------------------------
        # AUTO DOMAIN DETECTION
        # ---------------------------------------------
        auto_domain = infer_domain_from_data(df=df_ref, drift=drift_result["drift"])
        if auto_domain:
            st.session_state.domain = auto_domain
            st.info(f"📌 Auto-detected domain: **{auto_domain}**")
        # ===========================================================
        # CLEAN & SMOOTH DRIFT VISUALIZATION (BINNED LINE GRAPHS)
        # ===========================================================

        st.markdown("### 📈 Feature Drift Comparison")

        ref_proc = drift_result["ref_processed"]
        cur_proc = drift_result["cur_processed"]
        drift_vals = drift_result["drift"]

        # Only plot numeric columns that exist in both datasets
        features_to_plot = [
            c for c in ref_proc.columns
            if c in cur_proc.columns
               and pd.api.types.is_numeric_dtype(ref_proc[c])
        ]

        import numpy as np


        def smooth_bins(series, bins=100):
            """Aggregate values into N bins for smooth plotting."""
            try:
                ser = series.dropna().astype(float).values
                if len(ser) == 0:
                    return []
                # split into N bins
                binned = np.array_split(ser, bins)
                smoothed = [np.mean(b) if len(b) > 0 else None for b in binned]
                return smoothed
            except:
                return []


        cols_per_row = 2
        col_ptr = 0
        row = st.columns(cols_per_row)

        for feature in features_to_plot:

            ref_smooth = smooth_bins(ref_proc[feature])
            cur_smooth = smooth_bins(cur_proc[feature])

            fig = go.Figure()

            fig.add_trace(go.Scatter(
                y=ref_smooth,
                mode="lines",
                name="Reference (smoothed)",
                line=dict(width=3, color="#FF5733")
            ))

            fig.add_trace(go.Scatter(
                y=cur_smooth,
                mode="lines",
                name="Current (smoothed)",
                line=dict(width=3, color="#1E90FF")
            ))

            fig.update_layout(
                title=f"{feature} (drift: {drift_vals.get(feature, 0):.4f})",
                margin=dict(l=10, r=10, t=30, b=10),
                height=240,
                showlegend=True,
                legend=dict(orientation="h"),
            )

            with row[col_ptr]:
                st.plotly_chart(fig, use_container_width=True)

            col_ptr += 1
            if col_ptr == cols_per_row:
                col_ptr = 0
                row = st.columns(cols_per_row)
        # ======================================================================
        #                  PREMIUM DRIFT EXPLANATION ENGINE
        # ======================================================================

        drift_vals = drift_result.get("drift", {})

        # Sort highest → lowest
        sorted_items = sorted(drift_vals.items(), key=lambda x: x[1], reverse=True)


        # Severity thresholds
        def drift_severity(score):
            if score < 0.10:
                return "LOW", "🟢"
            elif score < 0.25:
                return "MEDIUM", "🟡"
            else:
                return "HIGH", "🔴"


        # Business meaning per feature (simple heuristic)
        def business_impact(feature, score):
            return (
                f"- This feature `{feature}` is behaving differently now.\n"
                f"- It may indicate a shift in customer behavior, system inputs, or real-world patterns.\n"
                f"- Because drift = **{score:.3f}**, the model may rely on outdated expectations."
            )


        # Fix suggestions
        def fix_suggestion(feature, score):
            if score < 0.10:
                return "- No urgent action needed."
            elif score < 0.25:
                return "- Monitor this feature closely.\n- Consider partial retraining if trend continues."
            else:
                return "- Retraining recommended.\n- Re-evaluate feature distributions.\n- Validate upstream data quality."


        # ======================================================================
        #                      DROPDOWN FOR EXPLANATION MODE
        # ======================================================================

        mode = st.selectbox(
            "Choose explanation type:",
            ["Layman Explanation", "Technical Explanation"]
        )

        # ======================================================================
        #                     LAYMAN EXPLANATION (PREMIUM)
        # ======================================================================

        if mode == "Layman Explanation":
            st.markdown("### Layman Explanation\nHere’s what changed:")

            for feature, score in sorted_items:
                sev, icon = drift_severity(score)

                with st.expander(f"{icon} {feature.replace('_', ' ').title()} — Changed by {score:.3f}"):
                    st.markdown(
                        f"""
                        **What it means (simple):**  
                        This part of your data has changed compared to the past.  
                        The model was expecting something different before.

                        **Severity:** {icon} **{sev} drift**
                        """
                    )

        # ======================================================================
        #                  TECHNICAL EXPLANATION (PREMIUM)
        # ======================================================================

        elif mode == "Technical Explanation":
            st.markdown("### Technical Explanation")

            for feature, score in sorted_items:
                sev, icon = drift_severity(score)

                with st.expander(f"{icon} {feature} — drift={score:.4f}"):
                    st.markdown(
                        f"""
                        - **Drift Score (PSI / distribution shift):** `{score:.4f}`
                        - **Severity:** {icon} **{sev}**

                        **PSI Interpretation:**
                        - 0.00–0.10: Minor  
                        - 0.10–0.25: Moderate  
                        - >0.25: High drift (retrain recommended)
                        """
                    )


        # ======================================================================
        #                 EXPORT CLEAN REPORT (TXT, DOCX, PDF)
        # ======================================================================

        st.markdown("### 📥 Download Full Drift Report")

        download_format = st.selectbox("Format:", ["TXT", "DOCX", "PDF"])

        report_text = "AI DRIFT RADAR REPORT\n=======================\n\n"

        for feature, score in sorted_items:
            sev, icon = drift_severity(score)
            report_text += f"""
        FEATURE: {feature}
        DRIFT: {score:.4f}
        SEVERITY: {sev}

        BUSINESS IMPACT:
        {business_impact(feature, score)}

        FIX SUGGESTIONS:
        {fix_suggestion(feature, score)}

        --------------------------------------------
        """

        # Build files
        from io import BytesIO
        from docx import Document
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.pagesizes import letter


        def make_txt():
            return report_text.encode("utf-8")


        def make_docx():
            doc = Document()
            for line in report_text.split("\n"):
                doc.add_paragraph(line)
            buf = BytesIO()
            doc.save(buf)
            return buf.getvalue()


        def make_pdf():
            buf = BytesIO()
            styles = getSampleStyleSheet()
            story = [Paragraph(line, styles["Normal"]) for line in report_text.split("\n")]
            SimpleDocTemplate(buf, pagesize=letter).build(story)
            return buf.getvalue()


        if download_format == "TXT":
            st.download_button("⬇ Download TXT", make_txt(), "drift_report.txt")

        elif download_format == "DOCX":
            st.download_button("⬇ Download DOCX", make_docx(), "drift_report.docx")

        else:
            st.download_button("⬇ Download PDF", make_pdf(), "drift_report.pdf")

# -----------------------------------------------
# PAGE: Model Monitor (metrics + embeddings + retrain logic)
# -----------------------------------------------
elif page == "Model Monitor":

    st.markdown("## 📈 Model Monitor")

    drift = st.session_state.last_drift
    if not drift:
        st.warning("⚠️ Compute drift first using 'Upload & Analyze'.")
    else:
        st.write("### Last Computed Drift")
        st.dataframe(pd.Series(drift).rename("score"))

    st.markdown("### 📄 Upload Model Metrics (metrics.json)")
    metrics_file = st.file_uploader("Upload metrics.json", type=["json"])

    st.markdown("### 🧬 Upload Embeddings (.npy)")
    ref_emb_file = st.file_uploader("Reference embeddings", type=["npy"])
    cur_emb_file = st.file_uploader("Current embeddings", type=["npy"])

    metrics = {}
    emb_shift = 0.0

    if metrics_file:
        try:
            metrics = json.load(metrics_file)
            st.json(metrics)
        except Exception:
            st.error("Invalid JSON.")

    if ref_emb_file and cur_emb_file:
        try:
            ref_emb = np.load(ref_emb_file)
            cur_emb = np.load(cur_emb_file)
            emb_shift = mean_cosine_embedding_shift(ref_emb, cur_emb)
            st.info(f"Embedding shift (cosine distance): {emb_shift:.4f}")
        except:
            st.error("Invalid embeddings.")

    st.markdown("---")
    st.markdown("## ⚖️ Auto-Retrain Evaluator")

    if drift or metrics or emb_shift:
        reasons = []
        total_drift = sum(drift.values()) if drift else 0

        if drift:
            if total_drift >= 0.25:
                reasons.append(f"Feature drift high: {round(total_drift, 4)} >= 0.25")

        for k, v in metrics.items():
            if v < 0.7:
                reasons.append(f"Metric '{k}' low: {v} < 0.7")

        if emb_shift >= 0.15:
            reasons.append(f"Embedding shift high: {emb_shift:.3f} >= 0.15")

        if reasons:
            st.error("⚠️ Auto-Retrain Recommended")
            for r in reasons:
                st.write("- " + r)
        else:
            st.success("✔ Model is stable. No retrain required.")

        payload = {
            "action": "trigger_retrain",
            "timestamp": datetime.utcnow().isoformat(),
            "reasons": reasons,
            "agg_feature_score": total_drift
        }

        st.markdown("### 📨 Webhook Payload")
        st.json(payload)

        st.code(f"""
curl -X POST https://your-system.com/webhook/retrain \\
  -H 'Content-Type: application/json' \\
  -d '{json.dumps(payload)}'
""")

    st.markdown("</div></div>", unsafe_allow_html=True)


# ---------------------------------------------------------
# PAGE: AI ASSISTANT (FINAL FIXED VERSION)
# ---------------------------------------------------------
elif page == "AI Assistant":

    st.markdown("## 🤖 AI Drift Assistant (Context-Aware + Domain-Aware)")

    # --- Initialize memory ---
    if "short_memory" not in st.session_state:
        st.session_state.short_memory = []
    if "domain" not in st.session_state:
        st.session_state.domain = ""


    def add_memory(role, content):
        st.session_state.short_memory.append({
            "role": role,
            "content": content,
            "time": datetime.utcnow().isoformat()
        })
        st.session_state.short_memory = st.session_state.short_memory[-12:]


    if not st.session_state.domain and "last_drift" in st.session_state:
        drift_data = st.session_state["last_drift"]

        # Check if drift_data is a dictionary and convert it to DataFrame if necessary
        if isinstance(drift_data, dict):
            # Convert dictionary to DataFrame
            drift_data = pd.DataFrame([drift_data])  # Convert dict to a DataFrame

        # Now call the infer_domain_from_data with the DataFrame
        auto = infer_domain_from_data(drift_data)
        if auto:
            st.session_state.domain = auto
            add_memory("assistant", f"Domain automatically inferred as {auto} based on drift patterns.")

    # --- Domain banner ---
    if st.session_state.domain:
        st.markdown(
            f"<div class='small' style='margin-bottom:8px;'>"
            f"Active Domain: <b>{st.session_state.domain}</b></div>",
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            "<div class='small' style='margin-bottom:8px;'>"
            "💡 Before starting, please type your domain (e.g., ecommerce, finance, or custom: telecom)."
            "</div>",
            unsafe_allow_html=True
        )

    # --- Render previous chat ---
    for turn in st.session_state.short_memory:
        cls = "chat-user" if turn["role"] == "user" else "chat-assistant"
        st.markdown(f"<div class='{cls}'>{turn['content']}</div>", unsafe_allow_html=True)

    placeholder_text = (
        "Type your domain to begin (e.g., ecommerce)"
        if not st.session_state.domain else
        "Ask anything about drift, embeddings, metrics, agents or retraining"
    )
    user_input = st.chat_input(placeholder_text)


    # ---------------------------
    #     UTILITY FUNCTIONS
    # ---------------------------
    def heuristic_is_ood(text):
        t = text.lower()
        ood = ["who is", "virat", "movie", "song", "weather", "capital", "president"]
        return any(x in t for x in ood)


    def wants_simplify(text):
        t = text.lower()
        return any(x in t for x in ["layman", "simple", "explain simply", "explain like"])


    # ---------------------------
    #     ON USER INPUT
    # ---------------------------
    if user_input:
        add_memory("user", user_input)

        # 1) Domain setup before anything else
        if not st.session_state.domain:
            detected = resolve_domain(user_input)
            if detected:
                st.session_state.domain = detected
                add_memory("assistant", f"✔ Domain set to {detected}. Now ask me drift/model questions.")
                st.rerun()
            else:
                add_memory("assistant",
                           "❗ I couldn't detect your domain. Please type ecommerce, finance, or custom: <name>.")
                st.rerun()

        # 2) User asked “explain in simple terms”
        if wants_simplify(user_input):
            # find last assistant message
            last_assistant = next(
                (m["content"] for m in reversed(st.session_state.short_memory) if m["role"] == "assistant"),
                None
            )
            if not last_assistant:
                add_memory("assistant", "I don’t have a previous explanation to simplify.")
                st.rerun()

            prompt = f"Simplify the following into easy layman terms:\n\n{last_assistant}"
            simple = groq_complete_sync(prompt, st.session_state.domain)
            add_memory("assistant", simple)
            st.rerun()

        # 3) Out-of-Domain queries blocked
        if heuristic_is_ood(user_input):
            add_memory(
                "assistant",
                "❌ This question is outside drift/model monitoring. Ask about drift, metrics, embeddings, or retraining."
            )
            st.rerun()

        # 4) Multi-agent analysis triggers
        triggers = ["drift", "psi", "embedding", "retrain", "degrade", "metrics", "analysis", "why"]
        needs_agents = any(t in user_input.lower() for t in triggers)

        domain = st.session_state.domain
        history = "\n".join([f"{m['role'].upper()}: {m['content']}" for m in st.session_state.short_memory[-8:]])
        drift_context = st.session_state.get("last_drift", {})

        # ---------------------------
        # MULTI-AGENT MODE
        # ---------------------------
        if needs_agents:
            st.info("🔎 Running multi-agent analysis...")

            da = agent_drift_analyst(history, domain)
            dq = agent_data_quality("reference", "current", domain)
            bi = agent_business_impact(history, domain)
            ra = agent_retrain_advisor(history, domain)
            oi = agent_ops_integration(history, domain)

            combined = f"""
Domain: {domain}
Latest drift: {drift_context}
User: {user_input}

Drift Analyst:
{da}

Data Quality:
{dq}

Business Impact:
{bi}

Retrain Advisor:
{ra}

Ops Integration:
{oi}

Synthesize the above into:
1) A two-line summary.
2) Four recommended actions.
3) A developer checklist.
"""
            final = groq_complete_sync(combined, domain)
            add_memory("assistant", final)
            st.rerun()

        # ---------------------------
        # NORMAL LLM REPLY
        # ---------------------------
        sys_prompt = f"""
You are AI Drift Radar Assistant.
Your rules:
- Only answer questions related to drift, embeddings, metrics, model degradation, retraining, or using this app.
- If outside topic → politely decline.
- Keep responses concise & structured.

Domain: {domain}
Recent context:
{history}
Last drift:
{drift_context}
"""

        reply = groq_complete_sync(f"{sys_prompt}\nUser: {user_input}", domain)
        add_memory("assistant", reply)
        st.rerun()

    # ---------------------------
    # CLEAR CHAT BUTTON
    # ---------------------------
    if st.button("🧹 Clear Chat"):
        st.session_state.short_memory = []
        st.session_state.domain = ""
        st.rerun()


# -----------------------------------------------
# PAGE: ABOUT
# -----------------------------------------------
elif page == "About":

    st.markdown("## ℹ️ About AI Drift Radar")
    st.write("""
    AI Drift Radar is a full-scale drift monitoring and model observability system built using:

    - **Python + Streamlit**
    - **Groq Llama 3.1-Instant** (LLM brain)
    - **Multi-Agent Architecture**
    - **PSI + Categorical Delta drift computation**
    - **Embedding drift via cosine distance**
    - **Metrics.json model performance integration**
    - **PDF/DOCX/TXT auto-report generator**
    - **Sample Data Engine** (10 domains)

    Designed for:
    - ML engineers  
    - Data scientists  
    - MLOps teams  
    - Architects  

    …to understand exactly *why* their models degrade and *what* to do about it.
    """)

    st.markdown("</div></div>", unsafe_allow_html=True)

